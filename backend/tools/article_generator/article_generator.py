"""Article Generator Workflow - Reusable agent for generating professional articles from documents.

This module provides a flexible article generation system that can assume different professional roles
to analyze and create articles from various document types (news, scientific papers, financial reports, etc.).
"""

import logging
import hashlib
import json
from typing import List, Dict, Union, Literal
from pathlib import Path
from datetime import datetime, timedelta
from pydantic import BaseModel, Field
from langchain_core.documents import Document
from tools.article_generator.article_prompts import (
    ArticleRole,
    get_role_context,
    generate_article_prompt,
    generate_digest_prompt,
    generate_qa_prompt
)
from tools.rag import AIChatClass
from langgraph.graph import StateGraph, START, END
from langgraph.prebuilt import ToolNode, tools_condition

from tools.rag.base_workflow import BaseWorkFlow, RAGState
from models.anthropic.token_counter import count_tokens_anthropic
from models.anthropic.usage_tracker import USAGE_TRACKER_INSTANCE as tracker


class ArticleStructure(BaseModel):
    """Structured output for generated articles."""
    
    title: str = Field(
        description="A compelling article title"
    )
    introduction: str = Field(
        description="Opening paragraph(s) that engage the reader"
    )
    main_sections: List[Dict[str, str]] = Field(
        description="List of section dictionaries with 'heading' and 'content' keys",
        min_items=2
    )
    conclusion: str = Field(
        description="Closing thoughts and actionable takeaways"
    )
    key_insights: List[str] = Field(
        description="3-5 main insights or key points from the article",
        min_items=3,
        max_items=5
    )
    tags: str = Field(
        description="Comma-separated relevant topic tags/keywords"
    )


class ArticleGenerator(BaseWorkFlow):
    """Generates professional articles from documents using configurable expert roles.
    
    This class extends BaseWorkflow to provide:
    - Configurable professional roles (financial advisor, journalist, researcher, etc.)
    - Document loading and analysis from multiple sources
    - Structured article generation with role-specific perspectives
    - Multi-article digest creation
    - Interactive Q&A about document content
    - Custom role definition support
    
    Example:
        generator = ArticleGenerator()
        generator.initialize(
            documents=['report.pdf', 'https://article.com'],
            role=ArticleRole.FINANCIAL_ADVISOR,
            article_style="formal",
            target_audience="investment professionals"
        )
        articles = generator.get_articles()
        digest = generator.get_digest()
    """
    
    def __init__(self, llm_model: AIChatClass = None, cache_dir: Union[str, Path] = None):
        """
        Initialize the article generator workflow.
        
        Args:
            llm_model: Optional LLM model. Defaults to llm_basic from models module
            cache_dir: Optional cache directory. Defaults to '.article_cache' in workspace root
        """
        super().__init__(llm_model)
        self._articles: List[Dict] = []
        self._role: ArticleRole = ArticleRole.JOURNALIST
        self._role_context: str = ""
        self._article_style: str = "formal"
        self._target_audience: str = "general audience"
        self._article_length: str = "medium"
        self._focus_areas: str = None
        self._custom_role_description: str = None
        self._consolidate_docs: bool = False
        self._cache_dir: Path = Path(cache_dir) if cache_dir else Path(".article_cache")
        self._cache_enabled: bool = True
        self._cache_expiry_days: int = 365
        self._cache_key: str = None
        self._total_tokens_used: int = 0
    
    def initialize(self,
                   documents: Union[List[str], List[Path]],
                   role: Union[ArticleRole, str] = ArticleRole.JOURNALIST,
                   custom_role_description: str = None,
                   article_style: Literal["formal", "conversational", "technical", "persuasive"] = "formal",
                   target_audience: str = "general audience",
                   article_length: Literal["short", "medium", "long"] = "medium",
                   focus_areas: str = None,
                   chunk_size: int = 2000,
                   chunk_overlap: int = 200,
                   name: str = "ArticleRetriever",
                   description: str = "Retrieves relevant document content for article generation",
                   consolidate_docs: bool = False,
                   use_cache: bool = True,
                   cache_expiry_days: int = 30):
        """
        Initialize the article generator with documents and role configuration.
        
        Args:
            documents: List of document paths or URLs to analyze
            role: Professional role (ArticleRole enum or string). Use 'custom' for custom roles
            custom_role_description: Custom role description when role='custom' or ArticleRole.CUSTOM
            article_style: Writing style - "formal", "conversational", "technical", or "persuasive"
            target_audience: Description of intended audience
            article_length: Article length - "short", "medium", or "long"
            focus_areas: Optional specific areas to focus on in articles
            chunk_size: Size of document chunks
            chunk_overlap: Overlap between chunks
            name: Name for the retriever
            description: Description of the retriever
            consolidate_docs: Whether to consolidate documents before processing
            use_cache: Whether to use persistent caching for generated articles
            cache_expiry_days: Number of days before cache expires (default: 30)
            
        Raises:
            ValueError: If role is CUSTOM but no custom_role_description provided
        """
        try:
            # Convert string role to ArticleRole enum if needed
            if isinstance(role, str):
                try:
                    self._role = ArticleRole(role.lower())
                except ValueError:
                    # If not a valid enum value, treat as custom
                    self._role = ArticleRole.CUSTOM
                    if not custom_role_description:
                        custom_role_description = f"You are a professional {role}."
            else:
                self._role = role
            
            # Store configuration
            self._custom_role_description = custom_role_description
            self._article_style = article_style
            self._target_audience = target_audience
            self._article_length = article_length
            self._focus_areas = focus_areas
            self._consolidate_docs = consolidate_docs
            self._cache_enabled = use_cache
            self._cache_expiry_days = cache_expiry_days
            self._total_tokens_used = 0
            
            # Get role context
            self._role_context = get_role_context(self._role, custom_role_description)
            
            logging.info(f"Initializing ArticleGenerator with role: {self._role}")
            logging.info(f"Role context: {self._role_context[:100]}...")
            
            # Load documents using base class method
            self._load_documents(
                data_sources=documents,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                use_cache=True
            )
            
            # Setup retriever and tools
            self._setup_retriever_and_tools(
                name=name,
                description=description
            )
            
            # Generate cache key based on configuration and documents
            self._cache_key = self._generate_cache_key(documents)
            
            # Try to load from cache
            if self._cache_enabled:
                cached_articles = self._load_from_cache()
                if cached_articles:
                    self._articles = cached_articles
                    logging.info(f"Loaded {len(cached_articles)} articles from cache.")
            
            logging.info(f"Article Generator initialized successfully.")
            
        except Exception as e:
            logging.error(f"Error initializing article generator: {e}")
            raise
    
    def _generate_article_from_document(self, document: Document) -> Dict[str, any]:
        """
        Generate a professional article from a single document.
        
        Args:
            document: Document containing source content
            
        Returns:
            Dictionary with article structure and metadata (including token counts)
        """
        try:
            # Create the prompt
            prompt = generate_article_prompt(
                role_context=self._role_context,
                document_content=document.page_content,
                article_style=self._article_style,
                target_audience=self._target_audience,
                article_length=self._article_length,
                focus_areas=self._focus_areas
            )
            
            # Count input tokens
            input_tokens = count_tokens_anthropic(
                messages=[{"role": "user", "content": prompt}],
                system=""
            )

            tracker.track_input(messages=[{"role": "user", "content": prompt}], tools=[])
            self._total_tokens_used += input_tokens
            logging.info(f"Total tokens used: {self._total_tokens_used}")
            
            # Use structured output
            structured_llm = self._llm_model.with_structured_output(ArticleStructure)
            response: ArticleStructure = structured_llm.invoke([{"role": "user", "content": prompt}])
            
            # Count output tokens by converting structured response to text
            # Combine all response fields into a single content string
            output_content = f"""Title: {response.title}
Introduction: {response.introduction}
Main Sections: {json.dumps(response.main_sections)}
Conclusion: {response.conclusion}
Key Insights: {json.dumps(response.key_insights)}
Tags: {response.tags}"""
            
            output_tokens = count_tokens_anthropic(
                messages=[{"role": "assistant", "content": output_content}],
                system=""
            )

            self._total_tokens_used += output_tokens
            logging.info(f"Total tokens: {self._total_tokens_used}")
            tracker.track_output(output_content)
            
            return {
                "title": response.title,
                "introduction": response.introduction,
                "main_sections": response.main_sections,
                "conclusion": response.conclusion,
                "key_insights": response.key_insights,
                "tags": response.tags,
                "metadata": {
                    "source": document.metadata.get('source', ''),
                    "role": self._role.value,
                    "style": self._article_style,
                    "audience": self._target_audience,
                    "original_length": len(document.page_content),
                    "input_tokens": input_tokens,
                    "output_tokens": output_tokens,
                    "total_tokens": input_tokens + output_tokens
                }
            }
            
        except Exception as e:
            logging.error(f"Error generating article: {e}")
            return {
                "title": f"Error: {document.metadata.get('title', 'Untitled')}",
                "introduction": f"Error generating article: {str(e)}",
                "main_sections": [],
                "conclusion": "",
                "key_insights": [],
                "tags": "",
                "metadata": {
                    "source": document.metadata.get('source', ''),
                    "error": str(e)
                }
            }
    
    def _generate_cache_key(self, documents: Union[List[str], List[Path]]) -> str:
        """Generate a unique cache key based on documents and configuration.
        
        Args:
            documents: List of document paths or URLs
            
        Returns:
            MD5 hash string as cache key
        """
        # Create a string combining all relevant parameters
        key_components = [
            str(sorted([str(d) for d in documents])),
            self._role.value,
            str(self._custom_role_description),
            self._article_style,
            self._target_audience,
            self._article_length,
            str(self._focus_areas),
            str(self._consolidate_docs)
        ]
        key_string = "|".join(key_components)
        return hashlib.md5(key_string.encode()).hexdigest()
    
    def _get_cache_path(self) -> Path:
        """Get the cache file path for current configuration.
        
        Returns:
            Path to cache file
        """
        self._cache_dir.mkdir(parents=True, exist_ok=True)
        return self._cache_dir / f"articles_{self._cache_key}.json"
    
    def _load_from_cache(self) -> Union[List[Dict], None]:
        """Load articles from cache if available and not expired.
        
        Returns:
            List of cached articles or None if cache miss/expired
        """
        cache_path = self._get_cache_path()
        
        if not cache_path.exists():
            logging.debug(f"Cache miss: {cache_path} does not exist.")
            return None
        
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
            
            # Check expiry
            cached_time = datetime.fromisoformat(cache_data.get('timestamp', ''))
            expiry_time = cached_time + timedelta(days=self._cache_expiry_days)
            
            if datetime.now() > expiry_time:
                logging.info(f"Cache expired (created: {cached_time}).")
                cache_path.unlink()  # Delete expired cache
                return None
            
            logging.info(f"Cache hit! Articles generated on {cached_time}.")
            return cache_data.get('articles', [])
            
        except Exception as e:
            logging.warning(f"Error loading cache: {e}")
            return None
    
    def _save_to_cache(self, articles: List[Dict]):
        """Save articles to cache.
        
        Args:
            articles: List of article dictionaries to cache
        """
        if not self._cache_enabled:
            return
        
        cache_path = self._get_cache_path()
        
        try:
            cache_data = {
                'timestamp': datetime.now().isoformat(),
                'cache_key': self._cache_key,
                'config': {
                    'role': self._role.value,
                    'style': self._article_style,
                    'audience': self._target_audience,
                    'length': self._article_length,
                    'focus_areas': self._focus_areas,
                    'consolidate_docs': self._consolidate_docs
                },
                'articles': articles
            }
            
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, indent=2, ensure_ascii=False)
            
            logging.info(f"Cached {len(articles)} articles to {cache_path}.")
            
        except Exception as e:
            logging.warning(f"Error saving to cache: {e}")
    
    def clear_cache(self, all_caches: bool = False):
        """Clear cached articles.
        
        Args:
            all_caches: If True, clear all caches. If False, clear only current cache.
        """
        if all_caches:
            if self._cache_dir.exists():
                import shutil
                shutil.rmtree(self._cache_dir)
                logging.info(f"Cleared all caches in {self._cache_dir}.")
        else:
            cache_path = self._get_cache_path()
            if cache_path.exists():
                cache_path.unlink()
                logging.info(f"Cleared cache: {cache_path}.")
    
    def _generate_all_articles(self) -> List[Dict]:
        """Generate articles from all loaded documents.
        
        Returns:
            List of article dictionaries
        """
        articles = []
        
        if not self._documents:
            logging.warning("No documents available to generate articles from.")
            return articles
        
        if self._consolidate_docs:
            
            logging.info("Consolidating documents before article generation...")
            combined_content = "\n\n".join(doc.page_content for doc in self._documents)
            consolidated_doc = Document(
                page_content=combined_content,
                metadata={"source": "Consolidated Documents"}
            )

            logging.info("Generating consolidated article...")
            article = self._generate_article_from_document(consolidated_doc)
            articles.append(article)
            
            # Cache the generated articles
            if articles and self._cache_enabled:
                self._save_to_cache(articles)
            
            return articles
        
        # Group documents by source to avoid duplicates
        docs_by_source = {}
        for doc in self._documents:
            source = doc.metadata.get('source', '')
            if source and source not in docs_by_source:
                docs_by_source[source] = doc
        
        for i, (source, doc) in enumerate(docs_by_source.items(), 1):
            logging.info(f"Generating article {i}/{len(docs_by_source)} from: {source[:80]}...")
            article = self._generate_article_from_document(doc)
            articles.append(article)
        
        # Cache the generated articles
        if articles and self._cache_enabled:
            self._save_to_cache(articles)
        
        return articles
    
    def generate_digest(self, articles: List[Dict] = None, digest_focus: str = "key themes and insights") -> str:
        """
        Generate a comprehensive digest from multiple articles.
        
        Args:
            articles: Optional list of article dictionaries. If not provided, uses self._articles
            digest_focus: What to focus on in the digest
            
        Returns:
            A comprehensive digest in Markdown format
        """
        if not articles:
            articles = self._articles
        
        if not articles:
            return "No articles to create digest from."
        
        # Prepare articles summary
        articles_text = []
        for i, article in enumerate(articles, 1):
            try:
                sections_summary = "\n".join([
                    f"  - {section.get('heading', 'Section')}: {section.get('content', '')[:200]}..."
                    for section in article.get('main_sections', [])[:3]
                ])
                
                articles_text.append(f"""Article {i}: {article['title']}
Source: {article['metadata'].get('source', 'N/A')}
Tags: {article['tags']}
Introduction: {article['introduction'][:300]}...
Key Insights: {chr(10).join(['- ' + insight for insight in article['key_insights']])}
Main Sections:
{sections_summary}""")
            except Exception as e:
                logging.error(f"Error preparing article {i} for digest: {e}")
                continue
        
        combined_text = '\n---\n'.join(articles_text)
        
        try:
            prompt = generate_digest_prompt(
                role_context=self._role_context,
                articles_summary=combined_text,
                digest_focus=digest_focus
            )
            response = self._llm_model.invoke([{"role": "user", "content": prompt}])
            return response.content
        
        except Exception as e:
            logging.error(f"Error generating digest: {e}")
            return f"Error generating digest: {str(e)}"
    
    def format_article_as_markdown(self, article: Dict) -> str:
        """
        Format an article dictionary as readable Markdown.
        
        Args:
            article: Article dictionary
            
        Returns:
            Markdown-formatted article string
        """
        md = f"# {article['title']}\n\n"
        md += f"**Tags:** {article['tags']}\n\n"
        md += f"---\n\n"
        md += f"{article['introduction']}\n\n"
        
        for section in article.get('main_sections', []):
            md += f"## {section.get('heading', 'Section')}\n\n"
            md += f"{section.get('content', '')}\n\n"
        
        md += f"## Conclusion\n\n"
        md += f"{article['conclusion']}\n\n"
        
        md += f"### Key Insights\n\n"
        for insight in article.get('key_insights', []):
            md += f"- {insight}\n"
        
        md += f"\n---\n\n"
        md += f"*Source: {article['metadata'].get('source', 'N/A')}*\n"
        md += f"*Role: {article['metadata'].get('role', 'N/A')} | Style: {article['metadata'].get('style', 'N/A')} | Audience: {article['metadata'].get('audience', 'N/A')}*\n"
        
        return md
    
    def get_articles(self, force_regenerate: bool = False) -> List[Dict]:
        """Get all generated articles.
        
        Args:
            force_regenerate: If True, ignore cache and regenerate articles
        
        Returns:
            List of article dictionaries
        """
        if force_regenerate:
            logging.info("Force regenerate: clearing cache and generating new articles.")
            self.clear_cache()
            self._articles = []
        
        if not self._articles:
           self._articles = self._generate_all_articles()
           if not self._articles:
                raise ValueError("No articles available. Call initialize() first.")
        
        return self._articles
    
    def get_digest(self, digest_focus: str = "key themes and insights") -> str:
        """Generate and return a comprehensive digest.
        
        Args:
            digest_focus: What to focus on in the digest
            
        Returns:
            Formatted digest string
        """
        if not self._articles:
           self._generate_all_articles()
           if not self._articles:
                raise ValueError("No articles available. Call initialize() first.")
        
        return self.generate_digest(self._articles, digest_focus)
    
    def save_articles_to_file(self, output_path: Union[str, Path], format: Literal["markdown", "json", "docx"] = "markdown"):
        """
        Save all articles to a file.
        
        Args:
            output_path: Path to save the articles
            format: Output format - "markdown", "json", or "docx" (for LinkedIn)
        """
        import json
        
        output_path = Path(output_path)
        
        if format == "markdown":
            content = f"# Generated Articles\n\n"
            content += f"**Role:** {self._role.value}\n"
            content += f"**Total Articles:** {len(self._articles)}\n\n"
            content += "---\n\n"
            
            for i, article in enumerate(self._articles, 1):
                content += f"\n\n<!-- Article {i} -->\n\n"
                content += self.format_article_as_markdown(article)
                content += "\n\n---\n\n"
            
            output_path.write_text(content, encoding='utf-8')
        
        elif format == "json":
            output_path.write_text(
                json.dumps(self._articles, indent=2, ensure_ascii=False),
                encoding='utf-8'
            )
        
        elif format == "docx":
            self._save_articles_as_docx(output_path)
        
        logging.info(f"Articles saved to {output_path}")
    
    def _save_articles_as_docx(self, output_path: Path):
        """
        Save articles to a Word document formatted for LinkedIn articles.
        
        LinkedIn article format:
        - Clean, professional styling
        - Clear headings and sections
        - Easy to copy/paste directly into LinkedIn article editor
        
        Args:
            output_path: Path to save the DOCX file
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Add header info
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(f"Generated Articles - {self._role.value.replace('_', ' ').title()}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Spacing
        
        # Add each article
        for i, article in enumerate(self._articles, 1):
            if i > 1:
                # Page break between articles
                doc.add_page_break()
            
            # Article title
            title = doc.add_heading(article['title'], level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Tags
            tags_para = doc.add_paragraph()
            tags_run = tags_para.add_run(f"#{article['tags'].replace(', ', ' #').replace(',', ' #')}")
            tags_run.italic = True
            tags_run.font.color.rgb = RGBColor(0, 119, 181)  # LinkedIn blue
            tags_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing
            
            # Introduction
            intro = doc.add_paragraph(article['introduction'])
            intro.style = 'Normal'
            
            doc.add_paragraph()  # Spacing
            
            # Main sections
            for section in article.get('main_sections', []):
                # Section heading
                heading = doc.add_heading(section.get('heading', 'Section'), level=2)
                
                # Section content
                content = doc.add_paragraph(section.get('content', ''))
                content.style = 'Normal'
                
                doc.add_paragraph()  # Spacing between sections
            
            # Conclusion heading
            conclusion_heading = doc.add_heading('Conclusion', level=2)
            
            # Conclusion text
            conclusion = doc.add_paragraph(article['conclusion'])
            conclusion.style = 'Normal'
            
            doc.add_paragraph()  # Spacing
            
            # Key Insights
            insights_heading = doc.add_heading('Key Takeaways', level=2)
            
            for insight in article.get('key_insights', []):
                insight_para = doc.add_paragraph(insight, style='List Bullet')
            
            # Metadata footer (optional - can be removed before posting)
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer_run = footer.add_run(
                f"\n___\n"
                f"Source: {article['metadata'].get('source', 'N/A')}\n"
                f"Generated by: {article['metadata'].get('role', 'N/A')} | "
                f"Style: {article['metadata'].get('style', 'N/A')}"
            )
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            footer_run.italic = True
        
        # Save document
        doc.save(str(output_path))
        logging.info(f"DOCX file saved with {len(self._articles)} article(s) ready for LinkedIn")
    
    def assemble_decision_flow(self):
        """Assemble article-specific decision flow for the workflow graph.
        
        Creates a LangGraph workflow for chat-based Q&A about the articles and source documents.
        """
        retriever_tool = self._get_retriever_tool()
        if not retriever_tool:
            raise ValueError("Failed to get retriever tool. Check initialization logs for details.")
        
        decision_flow = StateGraph(RAGState)
        
        # Simple flow: query -> retrieve -> answer
        decision_flow.add_node("query_documents", self.generate_query_or_respond)
        decision_flow.add_node("retrieve", ToolNode([retriever_tool]))
        decision_flow.add_node("answer", self.generate_article_answer)
        
        decision_flow.add_edge(START, "query_documents")
        
        # Use tools_condition to decide if we need retrieval
        decision_flow.add_conditional_edges(
            "query_documents",
            tools_condition,
            {
                "tools": "retrieve",
                END: END,
            },
        )
        
        decision_flow.add_edge("retrieve", "answer")
        decision_flow.add_edge("answer", END)
        
        return decision_flow
    
    def generate_article_answer(self, state: RAGState):
        """Generate answer for queries using retrieved context and article insights.
        
        Args:
            state: Current RAG state with messages and metadata
            
        Returns:
            Updated state with role-specific answer
        """
        question = state["messages"][0].content
        context = state["messages"][-1].content
        
        # Include article insights in the context
        articles_summary = "\n".join([
            f"- {a['title']}: {', '.join(a['key_insights'][:3])}"
            for a in self._articles[:5]  # Top 5 articles
        ])
        
        prompt = generate_qa_prompt(
            role_context=self._role_context,
            question=question,
            context=context,
            article_summaries=articles_summary
        )
        
        response = self._llm_model.invoke([{"role": "user", "content": prompt}])
        return {"messages": [response]}
    
    def run_example(self):
        """Example showing how to use the article generator."""
        if not self._articles:
            raise ValueError("No articles available. Call initialize() first.")
        
        print("\n" + "="*80)
        print(f"ARTICLE GENERATOR - {self._role.value.upper().replace('_', ' ')}")
        print("="*80 + "\n")
        
        # Show digest
        print("📊 GENERATING COMPREHENSIVE DIGEST...\n")
        digest = self.get_digest()
        print(digest)
        print("\n" + "="*80 + "\n")
        
        # Show individual articles
        print("📝 GENERATED ARTICLES:\n")
        for i, article in enumerate(self._articles, 1):
            print(f"\n{'='*80}")
            print(self.format_article_as_markdown(article))
            print(f"{'='*80}\n")
        
        # Example query
        if self._workflow_graph:
            print("\n" + "="*80)
            print("EXAMPLE Q&A")
            print("="*80 + "\n")
            
            result = self.invoke_chat([{
                "role": "user",
                "content": "What are the main insights from these documents?"
            }])
            print("Answer:", result["messages"][-1].content)


# Example usage:
# generator = ArticleGenerator()
# generator.initialize(
#     documents=['report.pdf', 'https://article.com'],
#     role=ArticleRole.FINANCIAL_ADVISOR,
#     article_style="formal",
#     target_audience="investment professionals"
# )
# articles = generator.get_articles()
# generator.save_articles_to_file('output_articles.md')
# digest = generator.get_digest()
