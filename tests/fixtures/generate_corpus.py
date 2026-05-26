"""One-off generator for synthetic test fixtures.

Run once (outputs are committed); rerun only if you change the fixture content.

Usage:
    .venv/Scripts/python.exe tests/fixtures/generate_corpus.py
"""
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

OUT = Path(__file__).parent / "corpus"
OUT.mkdir(parents=True, exist_ok=True)


def make_pdf() -> None:
    path = OUT / "sample-rocs-module.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "Sample ROCS Module — Fictional Test Fixture")
    c.setFont("Helvetica", 11)
    body = [
        "This is a synthetic test fixture, not a real FedEx document.",
        "",
        "The Sample Routing Module governs how the Sample Engine assigns",
        "fictional stops to fictional vehicles. It uses heuristic SAMPLE-A",
        "by default and falls back to SAMPLE-B when stop windows conflict.",
        "",
        "Error code SAMPLE-001 indicates the engine could not reconcile",
        "stop windows for a given vehicle. To resolve, planners should",
        "verify the stop manifest in the Sample Manifest screen and rerun",
        "the routing step. Error code SAMPLE-002 indicates a missing",
        "vehicle capacity profile; planners must populate the profile in",
        "the Sample Fleet screen.",
        "",
        "This document exists solely to exercise PyPDFLoader in tests.",
    ]
    y = 690
    for line in body:
        c.drawString(72, y, line)
        y -= 16
    c.showPage()
    c.save()
    print(f"wrote {path}")


def make_docx() -> None:
    path = OUT / "sample-error-catalog.docx"
    doc = Document()
    doc.add_heading("Sample Error Catalog — Fictional Test Fixture", level=1)
    doc.add_paragraph(
        "This is a synthetic test fixture, not a real FedEx document."
    )
    doc.add_heading("Error SAMPLE-100", level=2)
    doc.add_paragraph(
        "Triggered when the Sample Discrepancy Resolver finds a mismatch "
        "between planned and actual stop counts. Recommended action: "
        "open the Sample Reconciliation panel and accept or override the "
        "discrepancy line item."
    )
    doc.add_heading("Error SAMPLE-200", level=2)
    doc.add_paragraph(
        "Indicates the Sample Heuristic timed out. Recommended action: "
        "reduce the planning horizon to a single shift and rerun."
    )
    doc.save(str(path))
    print(f"wrote {path}")


def main() -> None:
    make_pdf()
    make_docx()


if __name__ == "__main__":
    main()
